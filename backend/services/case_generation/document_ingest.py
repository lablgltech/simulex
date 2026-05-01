"""Извлечение plain text из PDF, DOCX, Markdown для пайплайна генерации кейса."""

from __future__ import annotations

import os
from pathlib import Path
from typing import List, Tuple

                             
MAX_FILE_BYTES = int(os.getenv("CASE_GEN_MAX_FILE_BYTES", str(25 * 1024 * 1024)))
DEFAULT_MAX_CHARS = int(os.getenv("CASE_GEN_MAX_EXTRACT_CHARS", "500000"))

_ALLOWED_EXT = {".pdf", ".md", ".markdown", ".docx"}


def _strip_frontmatter_md(text: str) -> str:
    t = text.strip()
    if not t.startswith("---"):
        return text
    parts = t.split("---", 2)
    if len(parts) >= 3 and parts[0] == "":
        return parts[2].strip()
    return text


def extract_plain_text(
    file_bytes: bytes,
    filename: str,
    *,
    max_chars: int | None = None,
) -> Tuple[str, List[str]]:
    """
    Вернуть (текст, предупреждения).
    Бросает ValueError при неподдерживаемом формате или слишком большом файле.
    """
    warnings: List[str] = []
    if not file_bytes:
        raise ValueError("Пустой файл")

    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError(f"Файл больше {MAX_FILE_BYTES // (1024 * 1024)} МБ")

    ext = Path(filename or "upload").suffix.lower()
    if ext == ".doc":
        raise ValueError("Формат .doc не поддерживается, используйте DOCX или PDF")

    limit = max_chars if max_chars is not None else DEFAULT_MAX_CHARS

    if ext not in _ALLOWED_EXT:
        try:
            raw = file_bytes.decode("utf-8")
            return _apply_max_chars(raw.strip(), warnings + ["Файл без известного расширения обработан как UTF-8"], limit)
        except UnicodeDecodeError as e:
            raise ValueError(
                f"Неподдерживаемое расширение: {ext or '(нет)'}. Допустимы: PDF, DOCX, MD."
            ) from e

    if ext in (".md", ".markdown"):
        text = file_bytes.decode("utf-8", errors="replace")
        text = _strip_frontmatter_md(text)
        return _apply_max_chars(text.strip(), warnings, limit)

    if ext == ".docx":
        try:
            from io import BytesIO

            from docx import Document

            doc = Document(BytesIO(file_bytes))
            chunks: List[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    chunks.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        chunks.append("\t".join(cells))
            text = "\n\n".join(chunks)
        except Exception as e:
            raise ValueError(f"Не удалось прочитать DOCX: {e}") from e
        if not text.strip():
            warnings.append("DOCX: извлечённый текст пуст")
        return _apply_max_chars(text.strip(), warnings, limit)

    if ext == ".pdf":
        try:
            from io import BytesIO

            from pypdf import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            parts: List[str] = []
            for page in reader.pages:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                if t.strip():
                    parts.append(t)
            text = "\n\n".join(parts)
        except Exception as e:
            raise ValueError(f"Не удалось прочитать PDF: {e}") from e
        if not text.strip():
            warnings.append("PDF: текст не извлечён (возможно скан без OCR)")
        return _apply_max_chars(text.strip(), warnings, limit)

    raise ValueError(f"Неподдерживаемый формат: {ext}")


def _apply_max_chars(text: str, warnings: List[str], limit: int) -> Tuple[str, List[str]]:
    if len(text) > limit:
        warnings.append(f"Текст обрезан до {limit} символов")
        text = text[:limit]
    return text, warnings


def normalize_extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()
